import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import subprocess
import re
import docx # Import the python-docx library

# Ensure python-docx is installed: pip install python-docx

VOICES = [
    # English (en_GB)
    ("en_GB: alan (low)", "piper/piper-voices/en/en_GB/alan/low/en_GB-alan-low.onnx"),
    ("en_GB: alan (medium)", "piper/piper-voices/en/en_GB/alan/medium/en_GB-alan-medium.onnx"),
    ("en_GB: alba (medium)", "piper/piper-voices/en/en_GB/alba/medium/en_GB-alba-medium.onnx"),
    ("en_GB: aru (medium)", "piper/piper-voices/en/en_GB/aru/medium/en_GB-aru-medium.onnx"),
    ("en_GB: cori (medium)", "piper/piper-voices/en/en_GB/cori/medium/en_GB-cori-medium.onnx"),
    ("en_GB: cori (high)", "piper/piper-voices/en/en_GB/cori/high/en_GB-cori-high.onnx"),
    ("en_GB: jenny_dioco (medium)", "piper/piper-voices/en/en_GB/jenny_dioco/medium/en_GB-jenny_dioco-medium.onnx"),
    ("en_GB: northern_english_male (medium)", "piper/piper-voices/en/en_GB/northern_english_male/medium/en_GB-northern_english_male-medium.onnx"),
    ("en_GB: semaine (medium)", "piper/piper-voices/en/en_GB/semaine/medium/en_GB-semaine-medium.onnx"),
    ("en_GB: southern_english_female (low)", "piper/piper-voices/en/en_GB/southern_english_female/low/en_GB-southern_english_female-low.onnx"),
    ("en_GB: vctk (medium)", "piper/piper-voices/en/en_GB/vctk/medium/en_GB-vctk-medium.onnx"),

    # English (en_US)
    ("en_US: amy (low)", "piper/piper-voices/en/en_US/amy/low/en_US-amy-low.onnx"),
    ("en_US: amy (medium)", "piper/piper-voices/en/en_US/amy/medium/en_US-amy-medium.onnx"),
    ("en_US: arctic (medium)", "piper/piper-voices/en/en_US/arctic/medium/en_US-arctic-medium.onnx"),
    ("en_US: bryce (medium)", "piper/piper-voices/en/en_US/bryce/medium/en_US-bryce-medium.onnx"),
    ("en_US: danny (low)", "piper/piper-voices/en/en_US/danny/low/en_US-danny-low.onnx"),
    ("en_US: hfc_female (medium)", "piper/piper-voices/en/en_US/hfc_female/medium/en_US-hfc_female-medium.onnx"),
    ("en_US: hfc_male (medium)", "piper/piper-voices/en/en_US/hfc_male/medium/en_US-hfc_male-medium.onnx"),
    ("en_US: joe (medium)", "piper/piper-voices/en/en_US/joe/medium/en_US-joe-medium.onnx"),
    ("en_US: john (medium)", "piper/piper-voices/en/en_US/john/medium/en_US-john-medium.onnx"),
    ("en_US: kathleen (low)", "piper/piper-voices/en/en_US/kathleen/low/en_US-kathleen-low.onnx"),
    ("en_US: kristin (medium)", "piper/piper-voices/en/en_US/kristin/medium/en_US-kristin-medium.onnx"),
    ("en_US: kusal (medium)", "piper/piper-voices/en/en_US/kusal/medium/en_US-kusal-medium.onnx"),
    ("en_US: l2arctic (medium)", "piper/piper-voices/en/en_US/l2arctic/medium/en_US-l2arctic-medium.onnx"),
    ("en_US: lessac (low)", "piper/piper-voices/en/en_US/lessac/low/en_US-lessac-low.onnx"),
    ("en_US: lessac (medium)", "piper/piper-voices/en/en_US/lessac/medium/en_US-lessac-medium.onnx"),
    ("en_US: lessac (high)", "piper/piper-voices/en/en_US/lessac/high/en_US-lessac-high.onnx"),
    ("en_US: libritts (high)", "piper/piper-voices/en/en_US/libritts/high/en_US-libritts-high.onnx"),
    ("en_US: libritts_r (medium)", "piper/piper-voices/en/en_US/libritts_r/medium/en_US-libritts_r-medium.onnx"),
    ("en_US: ljspeech (medium)", "piper/piper-voices/en/en_US/ljspeech/medium/en_US-ljspeech-medium.onnx"),
    ("en_US: ljspeech (high)", "piper/piper-voices/en/en_US/ljspeech/high/en_US-ljspeech-high.onnx"),
    ("en_US: norman (medium)", "piper/piper-voices/en/en_US/norman/medium/en_US-norman-medium.onnx"),
    ("en_US: reza_ibrahim (medium)", "piper/piper-voices/en/en_US/reza_ibrahim/medium/en_US-reza_ibrahim-medium.onnx"),
    ("en_US: ryan (low)", "piper/piper-voices/en/en_US/ryan/low/en_US-ryan-low.onnx"),
    ("en_US: ryan (medium)", "piper/piper-voices/en/en_US/ryan/medium/en_US-ryan-medium.onnx"),
    ("en_US: ryan (high)", "piper/piper-voices/en/en_US/ryan/high/en_US-ryan-high.onnx"),
    ("en_US: sam (medium)", "piper/piper-voices/en/en_US/sam/medium/en_US-sam-medium.onnx"),

    # Spanish (es_ES/es_MX)
    ("es_ES: carlfm (x_low)", "piper/piper-voices/es/es_ES/carlfm/x_low/es_ES-carlfm-x_low.onnx"),
    ("es_ES: davefx (medium)", "piper/piper-voices/es/es_ES/davefx/medium/es_ES-davefx-medium.onnx"),
    ("es_ES: mls_10246 (low)", "piper/piper-voices/es/es_ES/mls_10246/low/es_ES-mls_10246-low.onnx"),
    ("es_ES: mls_9972 (low)", "piper/piper-voices/es/es_ES/mls_9972/low/es_ES-mls_9972-low.onnx"),
    ("es_ES: sharvard (medium)", "piper/piper-voices/es/es_ES/sharvard/medium/es_ES-sharvard-medium.onnx"),
    ("es_MX: ald (medium)", "piper/piper-voices/es/es_MX/ald/medium/es_MX-ald-medium.onnx"),
    ("es_MX: claude (high)", "piper/piper-voices/es/es_MX/claude/high/es_MX-claude-high.onnx"),

    # Portuguese (pt_BR/pt_PT)
    ("pt_BR: cadu (medium)", "piper/piper-voices/pt/pt_BR/cadu/medium/pt_BR-cadu-medium.onnx"),
    ("pt_BR: edresson (low)", "piper/piper-voices/pt/pt_BR/edresson/low/pt_BR-edresson-low.onnx"),
    ("pt_BR: faber (medium)", "piper/piper-voices/pt/pt_BR/faber/medium/pt_BR-faber-medium.onnx"),
    ("pt_BR: jeff (medium)", "piper/piper-voices/pt/pt_BR/jeff/medium/pt_BR-jeff-medium.onnx"),
    ("pt_PT: tugao (medium)", "piper/piper-voices/pt/pt_PT/tugao/medium/pt_PT-tugao-medium.onnx"),
]


WORDS_TO_REPLACE = [
    (" the hell ", " the heck "),
    (" the hell! ", " the heck! "),
    (" the hell? ", " the heck? "),

    (" the fuck ", " the heck "),
    (" the fuck! ", " the heck! "),
    (" the fuck? ", " the heck? "),
    (" the fucking hell ", " the heck "),

    (" fuck with ", " mess with "),
    (" fucking with ", " messing with "),
    (" fucking around ", " messing around "),

    (" fuck up ", " mess up "),
    (" fuckup ", " messup "),
    (" fuckups ", " messups "),
    (" fucker ", " mongrel "),
    (" fucker! ", " mongrel! "),
    (" fuckers ", " mongrels "),
    (" fuckers! ", " mongrels! "),
    (" fucking up ", " messing up "),
    (" Fuck! ", " Fetch! "),

    (" motherfucker ", " mongrel "),
    (" motherfucker! ", " mongrel! "),
    (" motherfuckers ", " mongrels "),
    (" motherfuckers! ", " mongrels! "),
    (" motherfucking ", " stupid "),

    (" god damn it ", " darnit "),
    (" goddamn it ", " darnit "),
    (" god damnit ", " darnit "),
    (" god damn it ", " darnit "),
    (" goddamn ", " darn "),
    (" goddamn! ", " darn! "),
    (" god damn ", " darn "),
    (" Damn! ", " Fetch! "),

    (" bullshit ", " nonsense "),
    (" bullshit! ", " nonsense! "),
    (" horsehit ", " nonsense "),
    (" horsehit! ", " nonsense! "),
    (" piece of shit ", " piece of trash "),
    (" little shit ", " scoundrel "),
    (" shitty ", " trashy "),
    (" shithole ", " trash heap "),
    (" shitless ", " "),
    (" Shit! ", " Fetch! "),
    (" shithole ", " trash heap "),

    (" crappy ", " trashy "),
    (" bullcrap ", " nonsense "),
    (" horsecrap ", " nonsense "),
    (" crapton ", " ton "),
    (" crap ton ", " ton "),
    (" cut the crap ", " cut the nonsense "),
    (" piece of crap ", " piece of trash "),
    (" Crap! ", " Fetch! "),

    (" bastard ", " mongrel "),
    (" Bastard! ", " Mongrel! "),
    (" bastards ", " mongrels "),
    (" Bastards! ", " Mongrels! "),

    (" retarded ", " idiotic "),
    (" retard ", " idiot "),
    (" retard! ", " idiot! "),
    (" retards ", " idiots "),
    (" retards! ", " idiots! "),

    (" bitchy ", " annoying "),
    (" bitching ", " whining "),
    (" Bitch! ", " Shrew! "),

    (" asshole ", " jerk "),
    (" asshole! ", " jerk! "),
    (" assholes ", " jerks "),
    (" assholes! ", " jerks! "),
    (" dumbass ", " idiot "),
    (" Dumbass! ", " Idiot! "),
    (" dumbasses ", " idiots "),
    (" Dumbasses! ", " Idiots! "),
    # Add more pairs as needed
]

# Moved get_project_root outside the class if it's logically tied to the GUI instance
# If it's a utility, it could remain outside. Let's keep it outside for now
# as it doesn't rely on instance state.
def get_project_root():
    """Finds the project root directory assuming it's named 'glowing-umbrella'."""
    path = os.path.abspath(os.getcwd())
    while True:
        if os.path.basename(path).lower() == "glowing-umbrella": # Use lower() for robustness
            return path
        parent = os.path.dirname(path)
        if parent == path:
            # Fallback: If not in a specific project structure, use current directory
            print("Warning: Could not find 'glowing-umbrella' project root. Using current directory.")
            return os.getcwd()
        path = parent


class PiperTTSGUI:
    def __init__(self, root):
        self.root = root
        root.title("Glowing Umbrella - Piper TTS Configuration")

        # Instance variables to hold file paths using StringVars for easy updating in Entry widgets
        self.input_folder_path = tk.StringVar() # For Text File Conversion
        self.output_folder_path = tk.StringVar() # For Text File Conversion Output
        self.text_file_path = tk.StringVar() # For Text File Split/Replace
        self.word_doc_path = tk.StringVar() # For Word File Split/Replace

        self.create_widgets()
        self.setup_validation()


    def create_widgets(self):
        # Use a frame for better layout management if needed, but grid is fine for this
        # main_frame = ttk.Frame(self.root, padding="10")
        # main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # --- Text File Conversion Section ---
        ttk.Label(self.root, text="Input Text Folder:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.input_path_entry = ttk.Entry(self.root, width=50, textvariable=self.input_folder_path)
        self.input_path_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(self.root, text="Browse...", command=self.browse_input).grid(row=0, column=2, padx=5, sticky="w")

        ttk.Label(self.root, text="Output Audio Folder:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.output_path_entry = ttk.Entry(self.root, width=50, textvariable=self.output_folder_path)
        self.output_path_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(self.root, text="Browse...", command=self.browse_output).grid(row=1, column=2, padx=5, sticky="w")

        # Voice
        ttk.Label(self.root, text="Voice:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.voice_combo = ttk.Combobox(
            self.root,
            values=[v[0] for v in VOICES],  # Combine lang tag and name
            state="readonly" # Prevent typing in the combobox
        )
        self.voice_combo.current(0) # Set default selected voice
        self.voice_combo.grid(row=2, column=1, padx=5, pady=5, columnspan=2, sticky="ew") # Span 2 columns

        # Speed/Pause
        ttk.Label(self.root, text="Speed Scale (0.001-2.0):").grid(row=3, column=0, padx=5, pady=5, sticky="w")
        self.speed_entry = ttk.Entry(self.root)
        self.speed_entry.insert(0, "1.0") # Default value
        self.speed_entry.grid(row=3, column=1, padx=5, pady=5, sticky="ew")

        ttk.Label(self.root, text="Pause (0.001-2.0):").grid(row=4, column=0, padx=5, pady=5, sticky="w")
        self.pause_entry = ttk.Entry(self.root)
        self.pause_entry.insert(0, "0.2") # Default value
        self.pause_entry.grid(row=4, column=1, padx=5, pady=5, sticky="ew")

        # Convert Button for Text Files
        ttk.Button(self.root, text="Convert Text Files to Audio", command=self.convert_files).grid(row=5, column=0, columnspan=3, pady=10)

        # Add a separator
        ttk.Separator(self.root, orient='horizontal').grid(row=6, column=0, columnspan=3, sticky="ew", pady=10)


        # --- Text File Utilities Section ---
        ttk.Label(self.root, text="Text File Utilities:").grid(row=7, column=0, padx=5, pady=5, sticky="w", columnspan=3)

        ttk.Label(self.root, text="Text File:").grid(row=8, column=0, padx=5, pady=5, sticky="w")
        self.txt_path_entry = ttk.Entry(self.root, width=50, textvariable=self.text_file_path)
        self.txt_path_entry.grid(row=8, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(self.root, text="Browse...", command=self.browse_txt).grid(row=8, column=2, padx=5, sticky="w")

        ttk.Label(self.root, text="Words per Split Text File:").grid(row=9, column=0, padx=5, pady=5, sticky="w")
        self.words_entry = ttk.Entry(self.root)
        self.words_entry.insert(0, "3700") # Default value
        self.words_entry.grid(row=9, column=1, padx=5, pady=5, sticky='ew')

        ttk.Button(self.root, text="Split Text File", command=self.split_text_file).grid(row=10, column=0, pady=10)
        ttk.Button(self.root, text="Replace Words in Text File", command=self.replace_words_in_text_file).grid(row=10, column=1, pady=10)


        # Add a separator
        ttk.Separator(self.root, orient='horizontal').grid(row=11, column=0, columnspan=3, sticky="ew", pady=10)


        # --- Word File Utilities Section ---
        ttk.Label(self.root, text="Word File Utilities (.docx):").grid(row=12, column=0, padx=5, pady=5, sticky="w", columnspan=3)

        ttk.Label(self.root, text="Word File:").grid(row=13, column=0, padx=5, pady=5, sticky="w")
        self.word_doc_path_entry = ttk.Entry(self.root, width=50, textvariable=self.word_doc_path)
        self.word_doc_path_entry.grid(row=13, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(self.root, text="Browse...", command=self.browse_word_doc).grid(row=13, column=2, padx=5, sticky="w")

        ttk.Label(self.root, text="Paragraphs per Split Word File:").grid(row=14, column=0, padx=5, pady=5, sticky="w")
        self.paragraphs_entry = ttk.Entry(self.root)
        self.paragraphs_entry.insert(0, "50") # Default split by 50 paragraphs
        self.paragraphs_entry.grid(row=14, column=1, padx=5, pady=5, sticky='ew')


        ttk.Button(self.root, text="Split Word File", command=self.split_word_doc).grid(row=15, column=0, pady=10)
        ttk.Button(self.root, text="Replace Words in Word File", command=self.replace_words_in_word_doc).grid(row=15, column=1, pady=10)


        # Add a separator
        ttk.Separator(self.root, orient='horizontal').grid(row=16, column=0, columnspan=3, sticky="ew", pady=10)

        # --- Other Utilities ---
        ttk.Button(self.root, text="Generate Voice Samples", command=self.generate_samples).grid(row=17, column=0, columnspan=3, pady=10)

        # Configure column weights so the entry fields expand
        self.root.grid_columnconfigure(1, weight=1)


    def setup_validation(self):
        # Register validation commands once
        vcmd_float = self.root.register(self.validate_float)
        vcmd_int = self.root.register(self.validate_int)

        self.speed_entry.config(validate="key", validatecommand=(vcmd_float, '%P'))
        self.pause_entry.config(validate="key", validatecommand=(vcmd_float, '%P'))
        self.words_entry.config(validate="key", validatecommand=(vcmd_int, '%P'))
        self.paragraphs_entry.config(validate="key", validatecommand=(vcmd_int, '%P'))


    def validate_float(self, value):
        """Validates if the input value is a float between 0.001 and 2.0."""
        try:
            if not value: return True # Allow empty during deletion
            f_value = float(value)
            return 0.001 <= f_value <= 2.0
        except ValueError:
            return False


    def validate_int(self, value):
        """Validates if the input value is an integer between 1 and 100000."""
        try:
            if not value: return True # Allow empty during deletion
            i_value = int(value)
            return 1 <= i_value <= 100000
        except ValueError:
            return False


    # --- Browse methods ---
    def browse_input(self):
        """Opens a dialog to select an input folder for text files."""
        path = filedialog.askdirectory(title="Select Input Folder (Text Files)")
        if path:
            self.input_folder_path.set(path)


    def browse_output(self):
        """Opens a dialog to select an output folder for audio files."""
        path = filedialog.askdirectory(title="Select Output Folder (Audio Files)")
        if path:
            self.output_folder_path.set(path)

    def browse_txt(self):
        """Opens a dialog to select a single text file for utilities."""
        path = filedialog.askopenfilename(
            title="Select Text File",
            filetypes=[("Text files", "*.txt"), ("All Files", "*.*")]
        )
        if path:
            self.text_file_path.set(path)

    def browse_word_doc(self):
        """Opens a dialog to select a single .docx file for utilities."""
        path = filedialog.askopenfilename(
            title="Select Word Document (.docx)",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if path:
            self.word_doc_path.set(path)


    # --- Text File Conversion Method ---
    def convert_files(self):
        """Converts text files from input folder to audio using Piper TTS."""
        try:
            # Get selected voice/model
            model_idx = self.voice_combo.current()
            if model_idx == -1:
                messagebox.showerror("Error", "Select a voice first.")
                return
            model_relpath = VOICES[model_idx][1]

            input_dir = self.input_folder_path.get()
            output_dir = self.output_folder_path.get()

            if not os.path.isdir(input_dir):
                 messagebox.showerror("Error", f"Input directory not found:\n{input_dir}")
                 return
            if not output_dir:
                 messagebox.showerror("Error", "Output directory is not selected.")
                 return
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)


            try:
                length_scale = float(self.speed_entry.get())
                sentence_silence = float(self.pause_entry.get())
                # Validation already handles basic float, but re-check range just in case
                if not (0.001 <= length_scale <= 2.0) or not (0.001 <= sentence_silence <= 2.0):
                     messagebox.showerror("Error", "Invalid speed or pause value range.")
                     return
            except ValueError:
                 messagebox.showerror("Error", "Invalid speed or pause value format.")
                 return


            # Find project root
            project_root = get_project_root() # Call the function

            # Build absolute model path
            model_path = os.path.join(project_root, model_relpath)
            if not os.path.isfile(model_path):
                messagebox.showerror("Error", f"Model file not found:\n{model_path}")
                return

            # Check for piper executable
            piper_executable = os.path.join(project_root, "piper_win", "piper.exe")
            if not os.path.isfile(piper_executable):
                 messagebox.showerror("Error", f"Piper executable not found:\n{piper_executable}")
                 return


            # Process each .txt file in input_dir
            files_converted = 0
            processed_files = 0
            for filename in os.listdir(input_dir):
                if not filename.lower().endswith(".txt"):
                    continue
                processed_files += 1
                base = os.path.splitext(filename)[0]
                input_path = os.path.join(input_dir, filename)
                output_path = os.path.join(output_dir, f"{base}.wav")

                # Construct the command using bash -c for robust path handling, especially on Windows+MSYS2
                # Use forward slashes in paths passed to bash commands
                command = (
                # Outer bash command uses double quotes
                # Inner double quotes around paths need to be escaped for bash (\")
                # The backslash used for bash escaping needs to be escaped for the Python string (\\")
                f'bash -c "MSYS2_ARG_CONV_EXCL=\\"*\\" cat \\"{input_path.replace("\\", "/")}\\" | '
                f'\\"{piper_executable.replace("\\", "/")}\\" '
                f'--model \\"{model_path.replace("\\", "/")}\\" --length_scale {length_scale:.2f} --sentence_silence {sentence_silence:.2f} '
                f'--output_file \\"{output_path.replace("\\", "/")}\\""'
                )

                print(f"Processing: {filename}")
                # print("COMMAND:", command) # Uncomment for debugging the command string
                # Use shell=True is necessary with bash -c and subprocess.run on Windows
                result = subprocess.run(command, shell=True, capture_output=True, text=True, encoding='utf-8')
                # print("STDOUT:", result.stdout) # Uncomment for debugging subprocess output
                # print("STDERR:", result.stderr) # Uncomment for debugging subprocess output

                if result.returncode == 0:
                    files_converted += 1
                else:
                    print(f"Failed to convert: {filename}")
                    print(f"Command: {command}")
                    print(f"Return Code: {result.returncode}")
                    print(f"Stderr: {result.stderr}")
                    messagebox.showwarning(f"Conversion Failed for {filename}", f"Piper command failed with return code {result.returncode}.\nSee console for details.")


            if processed_files == 0:
                 messagebox.showinfo("Info", f"No .txt files found in input directory:\n{input_dir}")
            else:
                messagebox.showinfo("Done", f"Attempted to convert {processed_files} files. Successfully converted {files_converted} files to audio.")
        except Exception as e:
            messagebox.showerror("Error", f"Conversion failed: {e}")

    # --- Text File Utility Methods ---
    def split_text_file(self):
        """Splits a text file by words."""
        txt_path = self.text_file_path.get()
        if not os.path.isfile(txt_path):
            messagebox.showerror("Error", "Selected text file does not exist.")
            return

        try:
            words_per_file = int(self.words_entry.get())
            if not 1 <= words_per_file <= 100000: # Validate range
                 messagebox.showerror("Error", "Words per file must be between 1 and 100000.")
                 return
        except ValueError:
            messagebox.showerror("Error", "Invalid value for words per file. Please enter a number.")
            return


        try:
            with open(txt_path, "r", encoding="utf-8") as f:
                words = f.read().split()
        except Exception as e:
            messagebox.showerror("Error", f"Could not read file:\n{e}")
            return

        if not words:
             messagebox.showinfo("Info", "The text file is empty or contains only whitespace.")
             return


        dir_name = os.path.dirname(txt_path)
        base_name = os.path.basename(txt_path)
        base, ext = os.path.splitext(base_name)
        if not ext: ext = ".txt" # Default to .txt if no extension


        file_idx = 1 # Start indexing from 1
        total_words = len(words)
        files_created = 0

        try:
            for i in range(0, total_words, words_per_file):
                chunk_words = words[i:i+words_per_file]
                # Generate output filename like original_part1.txt
                out_filename = f"{base}_part{file_idx}{ext}"
                out_path = os.path.join(dir_name, out_filename)

                # Simple overwrite prevention
                counter = 1
                original_out_path = out_path
                while os.path.exists(out_path):
                    out_filename = f"{base}_part{file_idx}_{counter}{ext}"
                    out_path = os.path.join(dir_name, out_filename)
                    counter += 1
                if counter > 1:
                     print(f"Warning: File '{original_out_path}' existed, saving as '{out_path}'")


                with open(out_path, "w", encoding="utf-8") as out_file:
                    out_file.write(" ".join(chunk_words))
                file_idx += 1
                files_created += 1
        except Exception as e:
            messagebox.showerror("Error", f"Could not write split file:\n{e}")
            return


        messagebox.showinfo("Success", f"Text file split into {files_created} files in directory:\n{dir_name}")


    def replace_words_in_text_file(self):
        """Replaces words/phrases in a text file."""
        txt_path = self.text_file_path.get()
        if not os.path.isfile(txt_path):
            messagebox.showerror("Error", "Selected text file does not exist.")
            return

        try:
            # Try multiple common encodings for text files
            text = None
            encodings_to_try = ['utf-8', 'cp1252', 'latin-1']
            for encoding in encodings_to_try:
                try:
                    with open(txt_path, "r", encoding=encoding) as f:
                        text = f.read()
                    print(f"Successfully read file with encoding: {encoding}")
                    break # Exit loop if successful
                except UnicodeDecodeError:
                    print(f"Failed to read with encoding: {encoding}")
                    continue
                except Exception as e:
                    # Handle other file reading errors
                    messagebox.showerror("Error", f"Error reading file {txt_path}:\n{e}")
                    return

            if text is None:
                 messagebox.showerror("Error", f"Could not decode text file {txt_path} using common encodings.")
                 return


            # Case-insensitive replacement using regex
            changes_made = False
            processed_text = text
            for old, new in WORDS_TO_REPLACE:
                # Use re.escape to handle special regex characters in the 'old' phrase
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                new_text = pattern.sub(new, processed_text)
                if new_text != processed_text:
                     changes_made = True
                     processed_text = new_text
                     # Optional: print which replacement happened
                     # print(f"Replaced occurrences of '{old}'")


            if changes_made:
                # Build output path: add _replaced before extension
                base, ext = os.path.splitext(txt_path)
                if not ext: ext = ".txt" # Default to .txt if no extension
                out_path = f"{base}_replaced{ext}"

                # Simple overwrite prevention
                counter = 1
                original_out_path = out_path
                while os.path.exists(out_path):
                    out_path = f"{base}_replaced_{counter}{ext}"
                    counter += 1
                if counter > 1:
                     print(f"Warning: Output file '{original_out_path}' existed, saving as '{out_path}'")


                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(processed_text)
                messagebox.showinfo("Success", f"File saved with replacements as:\n{out_path}")
            else:
                 messagebox.showinfo("Info", "No words/phrases from the list were found in the text file. No new file saved.")

        except Exception as e:
            messagebox.showerror("Error", f"Error processing text file:\n{e}")


    # --- Word File Utility Methods ---

    def split_word_doc(self):
        """
        Splits a Word .docx file into multiple files based on the number of paragraphs.
        Saves output files in the same directory as the source file.
        Preserves paragraph structure and attempts to preserve paragraph style (including font).
        """
        source_filepath = self.word_doc_path.get()
        if not os.path.isfile(source_filepath):
            messagebox.showerror("Error", "Selected Word file does not exist.")
            return

        # Warn if not a .docx, but allow attempting to open
        if not source_filepath.lower().endswith(".docx"):
            messagebox.showwarning("Warning", f"Selected file '{os.path.basename(source_filepath)}' does not have a .docx extension. Attempting to process anyway, but it might fail.")

        try:
            paragraphs_per_file_str = self.paragraphs_entry.get()
            # Use the validated integer directly
            if not paragraphs_per_file_str or int(paragraphs_per_file_str) < 1:
                 messagebox.showerror("Error", "Paragraphs per file must be at least 1.")
                 return
            paragraphs_per_file = int(paragraphs_per_file_str)

        except ValueError:
            messagebox.showerror("Error", "Invalid value for paragraphs per file. Please enter a number.")
            return


        output_dir = os.path.dirname(source_filepath)
        # Ensure output directory exists (it's the source dir, so it should, but being safe)
        os.makedirs(output_dir, exist_ok=True)

        try:
            document = docx.Document(source_filepath)
            paragraphs = document.paragraphs
            total_paragraphs = len(paragraphs)

            if total_paragraphs == 0:
                messagebox.showinfo("Info", f"The document {os.path.basename(source_filepath)} contains no paragraphs.")
                return

            base_name = os.path.splitext(os.path.basename(source_filepath))[0]
            created_files = []
            file_index = 1

            for i in range(0, total_paragraphs, paragraphs_per_file):
                new_document = docx.Document() # Creates a document with default styles
                chunk_paragraphs = paragraphs[i : i + paragraphs_per_file]

                for para in chunk_paragraphs:
                    # Add paragraph text, attempting to use the original paragraph style name.
                    # This preserves paragraph-level formatting (font, size, alignment etc.)
                    # but loses inline formatting within paragraphs (bold, italic etc.).
                    try:
                         # Check if the style name exists in the target document's styles
                         if para.style and para.style.name in new_document.styles:
                              new_document.add_paragraph(para.text, style=para.style.name)
                         else:
                              # Fallback to default style (usually 'Normal') if original style name is not found
                              new_document.add_paragraph(para.text)
                              if para.style: # Avoid warning for paragraphs without an explicit style
                                 print(f"Warning: Style '{para.style.name}' not found in new document for a paragraph. Used default style.")
                              # else:
                                 # This case is covered by default style usage, no warning needed

                    except Exception as add_para_error:
                         # Catch potential errors during paragraph addition (e.g., complex structures)
                         print(f"Warning: Failed to add paragraph text to new document: {add_para_error}. Adding as plain text.")
                         new_document.add_paragraph(para.text) # Fallback to adding plain text

                # Generate output filename like original_part1.docx
                out_filename = f"{base_name}_part{file_index}.docx"
                out_path = os.path.join(output_dir, out_filename)

                # Prevent overwriting if a file with the exact generated name already exists
                counter = 1
                original_out_name_for_msg = out_filename # Use this for user message if conflict occurs
                while os.path.exists(out_path):
                    out_filename = f"{base_name}_part{file_index}_{counter}.docx"
                    out_path = os.path.join(output_dir, out_filename)
                    counter += 1
                if counter > 1:
                     print(f"Warning: File '{original_out_name_for_msg}' existed, saving as '{os.path.basename(out_path)}'")

                new_document.save(out_path)
                created_files.append(out_path)
                file_index += 1

            if created_files:
                 messagebox.showinfo("Success", f"Word file split into {len(created_files)} files in directory:\n{output_dir}")
            else:
                 messagebox.showinfo("Info", "Word file is empty or too short to split into chunks of that size.")

        except docx.shared.exceptions.PackageNotFoundError:
            messagebox.showerror("Error", f"Could not open file {os.path.basename(source_filepath)}. It might not be a valid .docx file.")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred while splitting the Word file: {e}")


    def replace_words_in_word_doc(self):
        """
        Replaces words/phrases in a Word .docx file.
        Saves the modified file in the same directory as the source file.
        Preserves paragraph structure and attempts to preserve paragraph style (including font).
        Loses inline formatting within affected paragraphs.
        """
        source_filepath = self.word_doc_path.get()
        if not os.path.isfile(source_filepath):
            messagebox.showerror("Error", "Selected Word file does not exist.")
            return

        # Warn if not a .docx, but allow attempting to open
        if not source_filepath.lower().endswith(".docx"):
            messagebox.showwarning("Warning", f"Selected file '{os.path.basename(source_filepath)}' does not have a .docx extension. Attempting to process anyway, but it might fail.")


        output_dir = os.path.dirname(source_filepath)
        base_name = os.path.splitext(os.path.basename(source_filepath))[0]
        # Generate output file path: add _replaced before extension in the source directory
        output_filename = f"{base_name}_replaced.docx"
        output_filepath = os.path.join(output_dir, output_filename)

        # Prevent overwriting the original source file or a previously replaced file
        counter = 1
        original_output_name_for_msg = output_filename # Use this for user message if conflict occurs
        while os.path.exists(output_filepath):
            output_filename = f"{base_name}_replaced_{counter}.docx"
            output_filepath = os.path.join(output_dir, output_filename)
            counter += 1
        if counter > 1:
             print(f"Warning: Output file '{original_output_name_for_msg}' existed, saving as '{os.path.basename(output_filepath)}'")


        try:
            document = docx.Document(source_filepath)
            changes_made = False

            for old, new in WORDS_TO_REPLACE:
                # Create a case-insensitive regex pattern
                pattern = re.compile(re.escape(old), re.IGNORECASE)

                # Iterate through paragraphs to find and replace
                # Note: This approach iterates through the document and modifies it in place in memory.
                # If the document is huge, this might consume significant memory.
                for p in document.paragraphs:
                    original_text = p.text

                    # Perform the replacement using regex on the paragraph text
                    new_text = pattern.sub(new, original_text)

                    # If the text has changed, update the paragraph
                    if new_text != original_text:
                        # Store the original style object and alignment
                        original_style = p.style
                        original_alignment = p.alignment # Capture alignment too if style doesn't handle it

                        # Clear existing runs (this is where inline formatting is lost)
                        p.clear()

                        # Add the new text back as a single run
                        new_run = p.add_run(new_text)

                        # Reapply the original paragraph style and alignment
                        try:
                           p.style = original_style
                        except KeyError:
                            # Fallback if the style object is somehow invalid in this context
                            print(f"Warning: Could not reapply original style '{original_style.name}' to paragraph. Using default.")
                            p.style = document.styles['Normal'] # Fallback to Normal style

                        # Reapply alignment explicitly if needed (sometimes styles don't cover this completely)
                        if original_alignment is not None:
                             p.alignment = original_alignment

                        changes_made = True
                        # Optional: print which replacement happened in which paragraph
                        # print(f"Replaced '{old}' with '{new}' in a paragraph.")


            # --- Optional: Handle replacements in Tables ---
            # This adds significant complexity as tables contain cells, which contain paragraphs.
            # You would need nested loops: document.tables -> table.rows -> row.cells -> cell.paragraphs
            # for table in document.tables:
            #     for row in table.rows:
            #         for cell in row.cells:
            #             for p in cell.paragraphs:
            #                 original_text = p.text
            #                 new_text = pattern.sub(new, original_text)
            #                 if new_text != original_text:
            #                     original_style = p.style
            #                     original_alignment = p.alignment
            #                     p.clear()
            #                     p.add_run(new_text)
            #                     try: p.style = original_style
            #                     except KeyError: pass # Handle style not found
            #                     if original_alignment is not None: p.alignment = original_alignment
            #                     changes_made = True


            # --- Optional: Handle replacements in Headers and Footers ---
            # Requires iterating sections -> header/footer objects -> paragraphs within them.
            # for section in document.sections:
            #     # Headers
            #     if section.header is not None:
            #         for p in section.header.paragraphs:
            #              # ... apply replacement logic here ...
            #     # Footers
            #     if section.footer is not None:
            #         for p in section.footer.paragraphs:
            #              # ... apply replacement logic here ...


            if changes_made:
                document.save(output_filepath)
                messagebox.showinfo("Success", f"Word file saved with replacements as:\n{output_filepath}")
                return True
            else:
                 messagebox.showinfo("Info", f"No words/phrases from the list were found in '{os.path.basename(source_filepath)}'. No new file saved.")
                 return False

        except docx.shared.exceptions.PackageNotFoundError:
            messagebox.showerror("Error", f"Could not open file {os.path.basename(source_filepath)}. It might not be a valid .docx file.")
            return False
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred while replacing words in the Word file: {e}")
            return False


    # --- Generate Samples Method ---
    def generate_samples(self):
        """Generates sample audio files for English voices with various speed/pause settings."""
        try:
            project_root = get_project_root() # Call the function
            # Use os.path.join for robustness
            input_path = os.path.join(project_root, "input", "DO NOT DELETE", "DO NOT DELETE.txt")
            output_dir = os.path.join(project_root, "output", "samples")
            os.makedirs(output_dir, exist_ok=True) # Ensure output directory exists


            if not os.path.isfile(input_path):
                 messagebox.showerror("Error", f"Sample text file not found:\n{input_path}")
                 return

            # Check for piper executable
            # Use os.path.join for robustness
            piper_executable = os.path.join(project_root, "piper_win", "piper.exe")
            if not os.path.isfile(piper_executable):
                 messagebox.showerror("Error", f"Piper executable not found:\n{piper_executable}")
                 return


            length_scales = [0.6, 0.7, 0.8, 0.9, 1.00]
            sentence_silences = [0.05, 0.1, 0.15, 0.2, 0.25, 0.30]

            print("Generating voice samples...")
            for display_name, model_relpath in VOICES:
                # Only generate for English voices as previously decided/coded
                if not display_name.startswith("en_"):
                    continue
                model_path = os.path.join(project_root, model_relpath)
                if not os.path.isfile(model_path):
                     print(f"Skipping sample for missing model file: {model_relpath}")
                     continue

                base_model = os.path.splitext(os.path.basename(model_path))[0]
                for ls in length_scales:
                    for ss in sentence_silences:
                        output_file = f"{base_model}_ls{ls:.2f}_ss{ss:.2f}.wav"
                        output_path = os.path.join(output_dir, output_file)

                        # Construct the command using bash -c for robust path handling, especially on Windows+MSYS2
                        # Use forward slashes in paths passed to bash commands
                        command = (
                        # Outer bash command uses double quotes
                        # Inner double quotes around paths need to be escaped for bash (\")
                        # The backslash used for bash escaping needs to be escaped for the Python string (\\")
                        f'bash -c "MSYS2_ARG_CONV_EXCL=\\"*\\" cat \\"{input_path.replace("\\", "/")}\\" | '
                        f'\\"{piper_executable.replace("\\", "/")}\\" '
                        f'--model \\"{model_path.replace("\\", "/")}\\" --length_scale {length_scale:.2f} --sentence_silence {sentence_silence:.2f} '
                        f'--output_file \\"{output_path.replace("\\", "/")}\\""'
)

                        # print(f"Executing: {command}") # Uncomment for debugging the command string
                        # Use shell=True is necessary with bash -c and subprocess.run on Windows
                        result = subprocess.run(command, shell=True, capture_output=True, text=True, encoding='utf-8')
                        # print("STDOUT:", result.stdout) # Uncomment for debugging subprocess output
                        # print("STDERR:", result.stderr) # Uncomment for debugging subprocess output

                        if result.returncode != 0:
                            print(f"ERROR generating {output_path}:\nSTDOUT: {result.stdout}\nSTDERR: {result.stderr}")
                            # Optionally, break or show a messagebox here if an error stops the process
                        # else:
                        #     print(f"Generated: {output_path}") # Too verbose


            messagebox.showinfo("Done", f"Voice samples generated in directory:\n{output_dir}")
            print("Sample generation complete.")

        except Exception as e:
            messagebox.showerror("Error", f"Sample generation failed: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = PiperTTSGUI(root)
    root.mainloop()